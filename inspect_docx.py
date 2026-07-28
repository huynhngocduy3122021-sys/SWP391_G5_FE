from pathlib import Path
from docx import Document

source = Path(r"D:\Ki7\SWP391_G5_FE\Template2_SRD_Document2_working.docx")
doc = Document(source)

lines = []
for index, paragraph in enumerate(doc.paragraphs):
    text = paragraph.text.strip()
    if text:
        lines.append(f"P{index:04d} [{paragraph.style.name}]: {text}")

for table_index, table in enumerate(doc.tables):
    lines.append(f"\nTABLE {table_index}")
    for row_index, row in enumerate(table.rows):
        values = [" ".join(cell.text.split()) for cell in row.cells]
        lines.append(f"R{row_index:03d}: " + " || ".join(values))

output = Path(r"D:\Ki7\SWP391_G5_FE\document_content.txt")
output.write_text("\n".join(lines), encoding="utf-8")

print({
    "paragraphs": len(doc.paragraphs),
    "tables": len(doc.tables),
    "sections": len(doc.sections),
    "inline_shapes": len(doc.inline_shapes),
    "text_output": str(output),
})
